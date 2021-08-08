import datetime
import pprint
import sys
import re
import os

from collections import OrderedDict

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.shared import Pt

# 書式
#フォントサイズ、インデントの調整
normal_font_size = 10.5
middle_font_size = 12.0
large_font_size = 16.0

serihu_indent_chars_left = 7
serihu_indent_chars_right = 7
ch_name_indet_chars_left = 7
ch_name_indet_chars_right = 0

#段落の種類ごとのパラメータ
paragraph_params = {
    'title':
    {
        'bold': True,
        'font_size': large_font_size,
        'font_color': 'none',
        'left_indent': 0,
        'first_line_indent': 0,
        'right_indent': 0,
        'alignment': 'LEFT'
    },
    'sub_title':
    {
        'bold': True,
        'font_size': middle_font_size,
        'font_color': 'none',
        'left_indent': normal_font_size,
        'first_line_indent': 0,
        'right_indent': 1*middle_font_size,
        'alignment': 'LEFT'
    },
    'red_note':
    {
        'bold': False,
        'font_size': normal_font_size,
        'font_color': (210, 0, 0),
        'left_indent': 10*normal_font_size,
        'first_line_indent': 0,
        'right_indent': 0,
        'alignment': 'RIGHT'
    },
    'right_note':
    {
        'bold': False,
        'font_size': normal_font_size,
        'font_color': "none",
        'left_indent': 0,
        'first_line_indent': 0,
        'right_indent': normal_font_size,
        'alignment': 'RIGHT'
    },
    'blue_note':
    {
        'bold': False,
        'font_size': normal_font_size,
        'font_color': (0, 0, 210),
        'left_indent': 3*normal_font_size,
        'first_line_indent': 0,
        'right_indent': 3*normal_font_size,
        'alignment': 'LEFT'
    },    
    'hashira':
    {
        'bold': True,
        'font_size': normal_font_size,
        'font_color': 'none',
        'left_indent': 0,
        'first_line_indent': 0,
        'right_indent': 0,
        'alignment': 'LEFT'
    },
    'togaki':
    {
        'bold': False,
        'font_size': normal_font_size,
        'font_color': 'none',
        'left_indent': 0*normal_font_size,
        'first_line_indent': 0*normal_font_size,
        'right_indent': 0,
        'alignment': 'LEFT'
    },
    'character_name':
    {
       'bold': False,
        'font_size': normal_font_size,
        'font_color': 'none',
        'left_indent': (serihu_indent_chars_left + ch_name_indet_chars_left)*normal_font_size,
        'first_line_indent': 1*normal_font_size,
        'right_indent': (serihu_indent_chars_right + ch_name_indet_chars_right)*normal_font_size,
        'alignment': 'LEFT'
    },
    'character_serihu':
    {
        'bold': False,
        'font_size': normal_font_size,
        'font_color': 'none',
        'left_indent': serihu_indent_chars_left*normal_font_size,
        'first_line_indent': -1*normal_font_size,
        'right_indent': serihu_indent_chars_right*normal_font_size,
        'alignment': 'LEFT'
    },
    'kara':
    {
        'bold': False,
        'font_size': normal_font_size,
        'font_color': 'none',
        'left_indent': 0,
        'first_line_indent': 0,
        'right_indent': 0,
        'alignment': 'LEFT'
    },
    'plain':
    {
        'bold': False,
        'font_size': normal_font_size,
        'font_color': 'none',
        'left_indent': 0,
        'first_line_indent': 0,
        'right_indent': 0,
        'alignment': 'LEFT'
    }
}

html_header = '''
<!DOCTYPE html>
<html lang=\"ja\">
<head>
	<meta charset=\"UTF-8\">
	<title>preview</title>
	<style>
      html{
        font-size: 65.6%;
      }
      body {
        font-size: 1.4em;
        margin-left: 3em;
        font-family: serif;
      }
      article {
        width: 41em;
        margin-top: 2rem;
        margin-left: auto;
        margin-right: auto;
      }

      div.title{
        font-weight: bold; 
        font-size: 1.6em;
      }
      div.sub_title{
        font-weight: bold;
        font-size: 1.5em;
        padding-left: 1em;
      }
      div.red_note{
        color: #D20000;
        text-align:right;
      }
      div.right_note{
        text-align:right;
        padding-right: 1em;
      }
      div.blue_note{
        color: #0000D2;
        padding-left: 3em;
        padding-right: 3em;
      }
      div.hashira{
        font-weight: bold; 
      }
      div.togaki{
        text-indent: 0;
      }
      div.character_name{
        margin-top: 0em;
        padding-left: 15em;
      }
      div.character_serihu{
        margin-top: 0em;
        padding-left: 7em;
        padding-right: 7em;
        text-indent: -1em;
      }
      div.kara {
        margin-bottom: 1em;
      }
      div.plain {
        text-indent: 1em;
      }
  </style>
</head>
<body>
<article>
'''

html_footer = '''
</article>
</body>
</html>
'''


#段落の種類を判定
def identify(line):
    patterns = OrderedDict([
        ('.*', 'plain'),
        ('\s*\n', 'kara'),
        ('(.*)(「.*」)\s*\n', 'serihu'),
        ('(「.*」)\s*\n', 'character_serihu'),
        ('[0-9]+\.[\s　]+(.*)\n', 'character_name'),
        ('＠(.*)\n', 'togaki'),
        ('(■.*)\n', 'hashira'),
        ('>(.*)\n', 'blue_note'),
        ('→(.*)\n', 'right_note'),
        ('＃(.*)\n', 'red_note'),
        ('★(.*)\n', 'sub_title'),
        ('☆(.*)\n', 'title'),
    ])

    for k, v in patterns.items():
        if re.match(k, line):
            label = (v, k)
    return label

#ファイル読み込み
def file_reader(filename):
    basename, suffix = os.path.splitext(filename)
    if suffix == '.txt':
        with open(filename, encoding="utf-8") as f:
            lines = f.readlines()

    if suffix == '.docx':
        doc = docx.Document(filename)
        lines = []
        for paragraph in doc.paragraphs:
        #    lines.append(repr(paragraph.text) + '\n')
            lines.append(paragraph.text + '\n')
    return lines

#出力ファイル名を決める
def make_output_filename(filename):
    basename, suffix = os.path.splitext(filename)
    #now = datetime.datetime.now()
    #time_record = f"{now:%Y%m%d}"
    #output_filename =  basename + "_" + time_record + '_formatted' + '.docx'
    output_filename =  basename + '_formatted' + '.docx'
    return output_filename

def make_output_filename_html(filename):
    basename, suffix = os.path.splitext(filename)
    #now = datetime.datetime.now()
    #time_record = f"{now:%Y%m%d}"
    #output_filename =  basename + "_" + time_record + '_formatted' + '.docx'
    output_filename =  basename + '_formatted' + '.html'
    return output_filename

#出力時の段落に整形
def preprocessing(lines):
    status = 'status'
    line_count = 1
    formatted_lines = []
    for line in lines:
        type, reg = identify(line)
        # title, sub_title, red_note, right_note, blue_note, hashira, togaki, character_name, character_serihu
        if type in ['title', 'sub_title', 'red_note', 'right_note', 'blue_note', 'hashira', 'togaki', 'character_serihu']:
            text = re.match(reg, line).group(1)
            formatted_lines.append({'text': text, 'type': type})
        # kara
        elif type == 'kara':
            if status != 'serihu':
                formatted_lines.append({'text': ' ', 'type': 'kara'})
        #character_name
        elif type == 'character_name':
            text = str(line_count) + ".　" + re.match(reg, line).group(1)
            formatted_lines.append({'text': text, 'type': type})
            line_count += 1
        #serihu
        elif type == 'serihu':
            if status != 'serihu':
                formatted_lines.append({'text': ' ', 'type': 'kara'})
            m = re.match(reg, line)
  
            chracter_name = str(line_count) + '.　' + m.group(1)
            chracters_serihu = m.group(2)
            formatted_lines.append({'text': chracter_name, 'type': 'character_name'})
            formatted_lines.append({'text': chracters_serihu, 'type': 'character_serihu'})
            formatted_lines.append({'text': ' ', 'type': 'kara'})
            line_count += 1
        #plain
        else:
            text = line
            formatted_lines.append({'text': text.rstrip(), 'type': type})

        status = type
    return formatted_lines

#docx形式で出力
def paragraph_writer(text, params):
    paragraph = doc.add_paragraph(text)
    paragraph.runs[0].font.name = 'MS 明朝'


    # bold
    if params['bold'] == True:
        paragraph.runs[0].bold = True

    # font_size
    m = re.match('[0-9]+\.?[0-9]+', str(params['font_size']))
    if m:
        paragraph.runs[0].font.size = Pt(float(m.group(0)))

    # font_color
    if isinstance(params['font_color'], tuple):
        r, g, b, = params['font_color']
        paragraph.runs[0].font.color.rgb = RGBColor(r, g, b)

    # right-center-left
    if params['alignment'] == 'RIGHT':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # indentification
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Pt(params['left_indent'])
    paragraph_format.first_line_indent = Pt(
        params['first_line_indent'])
    paragraph_format.right_indent = Pt(params['right_indent'])

    # paragraph_space
    paragraph_format.space_before = Pt(0.0)
    paragraph_format.space_after = Pt(0.0)

#html形式で出力
def html_writer(text, type):
    return '<div class="' + type + '">' + text + '</div>' 


#メイン
doc = docx.Document()

filename = sys.argv[1]
output_filename = make_output_filename(filename)
output_filename_html = make_output_filename_html(filename)
source_lines = file_reader(filename)
preformatted_lines = preprocessing(source_lines)
pprint.pprint(preformatted_lines)

#docxファイル書き出し
for line in preformatted_lines:
    paragraph_writer(line['text'], paragraph_params[line['type']])

doc.save(output_filename)

#htmlファイル書き出し
html_doc = []
for line in preformatted_lines:
    html_doc.append(html_writer(line['text'], line['type']))

html = html_header + '\n'.join(html_doc) + html_footer

with open(output_filename_html, mode='w', encoding="utf-8") as f:
    f.write(html)
